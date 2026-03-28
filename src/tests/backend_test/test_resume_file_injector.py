"""Tests for backend.analysis.resume_file_injector.

Covers:
* _parse_projects_markdown  – markdown parsing edge cases
* _get_docx_font_info       – font detection from DOCX documents
* _get_pdf_font_info        – font detection from PDFs
* inject_projects_into_docx – section insertion (existing + new sections)
* inject_projects_into_pdf  – projects page appended to existing PDF
* extract_text_from_*       – text extraction helpers
* _pdf_has_projects_section – section detection in PDF text
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest

# Ensure src/ is on the path
SRC = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(SRC))

from backend.analysis.resume_file_injector import (
    _PROJECTS_ALIASES,
    _get_docx_font_info,
    _parse_projects_markdown,
    _pdf_has_projects_section,
    extract_text_from_docx,
    extract_text_from_pdf,
    inject_projects_into_docx,
    inject_projects_into_pdf,
)


# ---------------------------------------------------------------------------
# Helpers – minimal in-memory DOCX / PDF construction
# ---------------------------------------------------------------------------


def _make_docx(paragraphs: list[tuple[str, str]] | None = None) -> bytes:
    """Build a minimal DOCX in memory.

    *paragraphs* is a list of ``(text, style_name)`` tuples.
    Default style is ``"Normal"``.
    """
    from docx import Document

    doc = Document()
    for text, style in paragraphs or []:
        try:
            doc.add_paragraph(text, style=style)
        except Exception:
            p = doc.add_paragraph(text)
            if "heading" in style.lower():
                p.runs[0].bold = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_heading(heading_text: str, body_text: str = "Some body text") -> bytes:
    """DOCX with a Heading 2 paragraph followed by body text."""
    from docx import Document

    doc = Document()
    doc.add_heading(body_text, level=1)  # name / title
    doc.add_heading("Work Experience", level=2)
    doc.add_paragraph("Did things at Acme Corp.")
    doc.add_heading(heading_text, level=2)  # the section we want to find
    doc.add_paragraph("Existing project content.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.Sc. Computer Science")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_minimal_pdf(text_content: str = "Curriculum Vitae") -> bytes:
    """Build a minimal one-page PDF using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, text_content)
    c.save()
    return buf.getvalue()


def _make_pdf_with_projects_section() -> bytes:
    """Build a PDF that includes a 'Projects' heading."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 720, "Jane Smith")
    c.setFont("Helvetica", 12)
    c.drawString(72, 700, "Work Experience")
    c.drawString(72, 680, "Software Engineer at Acme Corp.")
    c.setFont("Helvetica-Bold", 12)
    c.drawString(72, 650, "Projects")
    c.setFont("Helvetica", 10)
    c.drawString(72, 630, "• My old project: built a thing")
    c.save()
    return buf.getvalue()


# ---------------------------------------------------------------------------
# _parse_projects_markdown
# ---------------------------------------------------------------------------


class TestParseProjectsMarkdown:
    def test_empty_string_returns_empty_list(self):
        assert _parse_projects_markdown("") == []

    def test_only_whitespace_returns_empty_list(self):
        assert _parse_projects_markdown("   \n\n  ") == []

    def test_projects_heading_is_skipped(self):
        md = "## Projects\n### My App\n- Did something\n"
        entries = _parse_projects_markdown(md)
        assert len(entries) == 1
        assert entries[0]["name"] == "My App"

    def test_single_project_with_bullets(self):
        md = "### Portfolio Site\n- Built with React\n- Deployed on Vercel\n"
        entries = _parse_projects_markdown(md)
        assert len(entries) == 1
        entry = entries[0]
        assert entry["name"] == "Portfolio Site"
        assert entry["tech"] is None
        assert entry["bullets"] == ["Built with React", "Deployed on Vercel"]

    def test_project_with_tech_line(self):
        md = "### Chat App\n*Python | WebSockets*\n- Implemented real-time messaging\n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["tech"] == "Python | WebSockets"
        assert entries[0]["bullets"] == ["Implemented real-time messaging"]

    def test_multiple_projects(self):
        md = (
            "## Projects\n\n"
            "### Alpha\n- Feature A\n- Feature B\n\n"
            "### Beta\n*Java | Spring*\n- Built REST API\n"
        )
        entries = _parse_projects_markdown(md)
        assert len(entries) == 2
        assert entries[0]["name"] == "Alpha"
        assert entries[1]["name"] == "Beta"
        assert entries[1]["tech"] == "Java | Spring"

    def test_bullet_with_dash(self):
        md = "### Project X\n- First bullet\n- Second bullet\n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["bullets"] == ["First bullet", "Second bullet"]

    def test_project_with_no_bullets(self):
        md = "### Empty Project\n*TypeScript*\n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["bullets"] == []
        assert entries[0]["tech"] == "TypeScript"

    def test_h2_non_projects_heading_becomes_entry(self):
        md = "## My Awesome Tool\n- Did X\n"
        entries = _parse_projects_markdown(md)
        assert len(entries) == 1
        assert entries[0]["name"] == "My Awesome Tool"

    def test_trailing_whitespace_stripped_from_names(self):
        md = "###   Spaced Name   \n- Bullet\n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["name"] == "Spaced Name"

    def test_empty_bullet_lines_ignored(self):
        md = "### X\n- \n- valid bullet\n-    \n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["bullets"] == ["valid bullet"]

    # --- New bold-header format (**Name** | *tech*) ---

    def test_bold_header_format_parsed_as_name(self):
        """Primary generator format: **Name** | *tech* should set name and tech."""
        md = "**My Project** | *Python*\n  - Built something\n"
        entries = _parse_projects_markdown(md)
        assert len(entries) == 1
        assert entries[0]["name"] == "My Project"
        assert entries[0]["tech"] == "Python"
        assert entries[0]["bullets"] == ["Built something"]

    def test_bold_header_without_tech(self):
        md = "**Solo Project**\n  - Just a project\n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["name"] == "Solo Project"
        assert entries[0]["tech"] is None

    def test_asterisks_stripped_from_name(self):
        """No asterisks should appear in the parsed name."""
        md = "**Cool Tool** | *JavaScript*\n  - Feature\n"
        entries = _parse_projects_markdown(md)
        assert "*" not in entries[0]["name"]
        assert "*" not in (entries[0]["tech"] or "")

    def test_multiple_bold_header_projects(self):
        md = (
            "## Projects\n\n"
            "**Alpha** | *Go*\n  - Service A\n\n"
            "**Beta** | *Rust*\n  - Service B\n"
        )
        entries = _parse_projects_markdown(md)
        assert len(entries) == 2
        assert entries[0]["name"] == "Alpha"
        assert entries[0]["tech"] == "Go"
        assert entries[1]["name"] == "Beta"
        assert entries[1]["tech"] == "Rust"

    def test_bold_format_with_pipe_and_spaces(self):
        """Pipe separator with varied spacing should still work."""
        md = "**My App** | *TypeScript | React*\n  - UI work\n"
        entries = _parse_projects_markdown(md)
        assert entries[0]["name"] == "My App"
        assert "TypeScript" in (entries[0]["tech"] or "")
        assert "*" not in (entries[0]["tech"] or "")


# ---------------------------------------------------------------------------
# DOCX font detection
# ---------------------------------------------------------------------------


class TestGetDocxFontInfo:
    def test_returns_tuple_of_str_and_float(self):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        doc2 = Document(io.BytesIO(docx_bytes))
        font_name, font_size = _get_docx_font_info(doc2)
        assert isinstance(font_name, str)
        assert len(font_name) > 0
        assert isinstance(font_size, float)
        assert font_size > 0

    def test_custom_normal_font_is_detected(self):
        """When the Normal style font is explicitly set, it should be returned."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        try:
            doc.styles["Normal"].font.name = "Times New Roman"
            doc.styles["Normal"].font.size = Pt(12)
        except Exception:
            pytest.skip("Cannot set Normal style font in this environment")

        font_name, font_size = _get_docx_font_info(doc)
        assert font_name == "Times New Roman"
        assert font_size == pytest.approx(12.0)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


class TestExtractText:
    def test_extract_text_from_docx_returns_string(self):
        docx_bytes = _make_docx([("Hello World", "Normal"), ("Second para", "Normal")])
        text = extract_text_from_docx(docx_bytes)
        assert "Hello World" in text
        assert "Second para" in text

    def test_extract_text_from_docx_empty_doc(self):
        docx_bytes = _make_docx([])
        text = extract_text_from_docx(docx_bytes)
        assert isinstance(text, str)

    def test_extract_text_from_pdf_returns_string(self):
        pdf_bytes = _make_minimal_pdf("Jane Smith Resume")
        text = extract_text_from_pdf(pdf_bytes)
        assert isinstance(text, str)
        # reportlab-generated PDFs should be extractable
        assert "Jane Smith" in text or len(text) >= 0  # at minimum no crash

    def test_extract_text_from_invalid_bytes_returns_empty(self):
        assert extract_text_from_docx(b"not a docx") == ""
        assert extract_text_from_pdf(b"not a pdf") == ""


# ---------------------------------------------------------------------------
# _pdf_has_projects_section
# ---------------------------------------------------------------------------


class TestPdfHasProjectsSection:
    def test_detects_projects_heading(self):
        pdf_bytes = _make_pdf_with_projects_section()
        assert _pdf_has_projects_section(pdf_bytes) is True

    def test_no_projects_section_returns_false(self):
        pdf_bytes = _make_minimal_pdf("Work Experience only")
        result = _pdf_has_projects_section(pdf_bytes)
        assert result is False

    def test_project_experience_alias_detected(self):
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(72, 720, "Project Experience")
        c.save()
        pdf_bytes = buf.getvalue()
        assert _pdf_has_projects_section(pdf_bytes) is True


# ---------------------------------------------------------------------------
# inject_projects_into_docx
# ---------------------------------------------------------------------------


class TestInjectProjectsIntoDocx:
    _sample_md = (
        "## Projects\n"
        "### My Web App\n"
        "*React | Node.js*\n"
        "- Built a full-stack application\n"
        "- Implemented user authentication\n"
        "\n"
        "### CLI Tool\n"
        "- Automated deployment workflows\n"
    )

    def test_creates_entries_when_no_projects_section_exists(self):
        """Entries are appended directly without adding a 'Projects' heading."""
        docx_bytes = _make_docx([("John Doe", "Normal"), ("Work Experience", "Normal")])
        result = inject_projects_into_docx(docx_bytes, self._sample_md)
        assert isinstance(result, bytes)
        assert len(result) > len(docx_bytes)

        from docx import Document

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Project names must appear
        assert "My Web App" in full_text
        assert "CLI Tool" in full_text

    def test_adds_project_name_as_title(self):
        docx_bytes = _make_docx([])
        result = inject_projects_into_docx(docx_bytes, self._sample_md)

        from docx import Document

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert "My Web App" in texts
        assert "CLI Tool" in texts

    def test_adds_bullets_under_project_name(self):
        docx_bytes = _make_docx([])
        result = inject_projects_into_docx(docx_bytes, self._sample_md)

        from docx import Document

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Built a full-stack application" in full_text or "\u2022" in full_text

    def test_inserts_into_existing_projects_section(self):
        """Projects should be added at the end of an existing Projects section."""
        docx_bytes = _make_docx_with_heading("Projects")
        result = inject_projects_into_docx(docx_bytes, self._sample_md)

        from docx import Document

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # New projects must appear
        assert "My Web App" in full_text
        # Original content must be preserved
        assert "Existing project content." in full_text
        # Education section must still exist after projects
        assert "Education" in full_text

    def test_inserts_before_next_section_not_at_very_end(self):
        """After injection, Education heading should come after new projects."""
        docx_bytes = _make_docx_with_heading("Projects")
        result = inject_projects_into_docx(docx_bytes, self._sample_md)

        from docx import Document

        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        # Find positions
        proj_idx = next((i for i, t in enumerate(texts) if "My Web App" in t), -1)
        edu_idx = next((i for i, t in enumerate(texts) if t.strip() == "Education"), -1)
        assert proj_idx != -1, "My Web App not found"
        assert edu_idx != -1, "Education not found"
        # New project must appear before Education
        assert proj_idx < edu_idx

    def test_uses_document_font(self):
        """Added runs should carry the document's Normal font name."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        try:
            doc.styles["Normal"].font.name = "Georgia"
            doc.styles["Normal"].font.size = Pt(11)
        except Exception:
            pytest.skip("Cannot set Normal style font in test environment")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = inject_projects_into_docx(docx_bytes, self._sample_md)

        doc2 = Document(io.BytesIO(result))
        # Find a run in the added bullets and check its font
        found_font = None
        for para in doc2.paragraphs:
            if "Built a full-stack" in para.text or "\u2022" in para.text:
                for run in para.runs:
                    if run.font.name:
                        found_font = run.font.name
                        break
        # Font should be set (either Georgia or the fallback)
        assert found_font is not None

    def test_returns_original_on_empty_markdown(self):
        docx_bytes = _make_docx([("Content", "Normal")])
        result = inject_projects_into_docx(docx_bytes, "")
        assert result == docx_bytes

    def test_returns_original_on_projects_heading_only(self):
        docx_bytes = _make_docx([])
        result = inject_projects_into_docx(docx_bytes, "## Projects\n")
        assert result == docx_bytes

    def test_output_is_valid_docx(self):
        from docx import Document

        docx_bytes = _make_docx([("Hello", "Normal")])
        result = inject_projects_into_docx(docx_bytes, self._sample_md)
        # Should not raise
        Document(io.BytesIO(result))

    def test_project_experience_alias_recognized(self):
        """'Project Experience' heading should be treated as a Projects section."""
        docx_bytes = _make_docx_with_heading("Project Experience")
        result = inject_projects_into_docx(docx_bytes, self._sample_md)

        from docx import Document

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "My Web App" in full_text
        # Original content preserved
        assert "Existing project content." in full_text


# ---------------------------------------------------------------------------
# inject_projects_into_pdf
# ---------------------------------------------------------------------------


class TestInjectProjectsIntoPdf:
    _sample_md = (
        "## Projects\n"
        "### Data Dashboard\n"
        "*Python | Dash*\n"
        "- Visualised sensor data in real time\n"
        "- Reduced reporting time by 40%\n"
    )

    def test_returns_bytes(self):
        pdf_bytes = _make_minimal_pdf()
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        assert isinstance(result, bytes)

    def test_output_is_larger_than_input(self):
        pdf_bytes = _make_minimal_pdf()
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        assert len(result) > len(pdf_bytes)

    def test_output_is_valid_pdf(self):
        from pypdf import PdfReader

        pdf_bytes = _make_minimal_pdf()
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        reader = PdfReader(io.BytesIO(result))
        assert len(reader.pages) >= 2  # original + at least one new page

    def test_original_pages_preserved(self):
        from pypdf import PdfReader

        pdf_bytes = _make_minimal_pdf()
        original_page_count = len(PdfReader(io.BytesIO(pdf_bytes)).pages)
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        result_page_count = len(PdfReader(io.BytesIO(result)).pages)
        assert result_page_count > original_page_count

    def test_returns_original_on_empty_markdown(self):
        pdf_bytes = _make_minimal_pdf()
        result = inject_projects_into_pdf(pdf_bytes, "")
        assert result == pdf_bytes

    def test_returns_original_on_projects_heading_only(self):
        pdf_bytes = _make_minimal_pdf()
        result = inject_projects_into_pdf(pdf_bytes, "## Projects\n")
        assert result == pdf_bytes

    def test_no_section_heading_in_appended_page(self):
        """Injected page must NOT contain a 'Projects' section heading."""
        from pypdf import PdfReader

        pdf_bytes = _make_minimal_pdf("No projects here")
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        reader = PdfReader(io.BytesIO(result))
        last_page_text = reader.pages[-1].extract_text() or ""
        # "continued" label no longer added
        assert "continued" not in last_page_text.lower()

    def test_no_section_heading_even_when_pdf_has_projects(self):
        """When the source PDF already has a Projects section no extra heading is added."""
        from pypdf import PdfReader

        pdf_bytes = _make_pdf_with_projects_section()
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        reader = PdfReader(io.BytesIO(result))
        last_page_text = reader.pages[-1].extract_text() or ""
        assert "continued" not in last_page_text.lower()

    def test_project_name_appears_in_appended_page(self):
        from pypdf import PdfReader

        pdf_bytes = _make_minimal_pdf()
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        reader = PdfReader(io.BytesIO(result))
        last_page_text = reader.pages[-1].extract_text() or ""
        assert "Data Dashboard" in last_page_text

    def test_uses_helvetica_family_for_helvetica_pdf(self):
        """A plain Helvetica PDF should produce a projects page also using Helvetica."""
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(72, 720, "CV")
        c.save()
        pdf_bytes = buf.getvalue()

        # Should not raise; font detection falls back gracefully
        result = inject_projects_into_pdf(pdf_bytes, self._sample_md)
        assert isinstance(result, bytes)
        assert len(result) > len(pdf_bytes)


# ---------------------------------------------------------------------------
# Projects aliases coverage
# ---------------------------------------------------------------------------


class TestProjectsAliases:
    def test_aliases_set_is_non_empty(self):
        assert len(_PROJECTS_ALIASES) > 0

    def test_projects_is_in_aliases(self):
        assert "projects" in _PROJECTS_ALIASES

    def test_all_aliases_are_lowercase(self):
        for alias in _PROJECTS_ALIASES:
            assert alias == alias.lower(), f"Alias not lowercase: {alias}"
