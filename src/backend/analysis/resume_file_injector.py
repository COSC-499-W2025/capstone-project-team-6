"""Utilities for injecting generated project entries into existing PDF or DOCX resumes.

Design
------
* **DOCX** – detect the document's default body font (from the Normal style or theme),
  then use that same font on every run added to the document.  If a Projects section
  already exists (matching a known alias list), new entries are inserted at the end of
  that section; otherwise the entries are appended directly at the end of the document
  with no extra section heading.

* **PDF** – detect the primary body font from the first page's resource dictionary and
  map it to the nearest built-in reportlab font.  Project entries are appended as a new
  page at the end without any section heading (in-page injection is not feasible without
  PyMuPDF).

In both cases no "Projects" section heading is ever created – entries (name, tech, and
bullets) are inserted as-is so the caller's document structure is preserved.
"""

from __future__ import annotations

import io
import logging
import re
from typing import List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Common heading aliases that indicate a projects section
# ---------------------------------------------------------------------------

_PROJECTS_ALIASES = frozenset(
    {
        "projects",
        "project experience",
        "personal projects",
        "side projects",
        "my projects",
        "software projects",
        "technical projects",
        "relevant projects",
        "selected projects",
    }
)


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract plain text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n".join(parts)
    except Exception:
        logger.exception("Failed to extract text from PDF")
        return ""


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception:
        logger.exception("Failed to extract text from DOCX")
        return ""


# ---------------------------------------------------------------------------
# Markdown → structured project entries
# ---------------------------------------------------------------------------


def _strip_md(text: str) -> str:
    """Remove markdown bold/italic asterisks and backticks from *text*."""
    # Remove ** and * markers; remove backticks
    text = re.sub(r"\*+", "", text)
    text = text.replace("`", "")
    return text.strip()


def _parse_projects_markdown(projects_markdown: str) -> List[dict]:
    """Parse resume-generator markdown into ``[{"name", "tech", "bullets"}]``.

    Handles the primary format emitted by ``_format_bundle``:

        **Project Name** | *Python*
          - Bullet one
          - Bullet two

    Also handles ``### Project Name`` and ``## Project Name`` heading styles.
    All markdown formatting characters (``*``, ``**``, backticks) are stripped
    from names and tech strings so they render cleanly in documents.
    """
    entries: list[dict] = []
    current: Optional[dict] = None

    for line in projects_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        # ---- Bold header: **Project Name** | *tech*  (primary generator format) ----
        if stripped.startswith("**") and not stripped.startswith("***"):
            if current:
                entries.append(current)
            # Strip the leading **
            rest = stripped[2:]  # "Project Name** | *Python*"  or  "Name**"
            if "**" in rest:
                name_raw, _, tech_raw = rest.partition("**")
                name = _strip_md(name_raw)
                tech_raw = tech_raw.strip().lstrip("|").strip()
                tech = _strip_md(tech_raw) if tech_raw else None
            else:
                name = _strip_md(rest)
                tech = None
            current = {"name": name, "tech": tech or None, "bullets": []}

        # ---- Heading: ### Project Name ----
        elif stripped.startswith("###"):
            if current:
                entries.append(current)
            current = {"name": _strip_md(stripped.lstrip("#")), "tech": None, "bullets": []}

        # ---- Heading: ## Project Name  (skip "Projects" section heading) ----
        elif stripped.startswith("##"):
            title = _strip_md(stripped.lstrip("#"))
            if title.lower() in _PROJECTS_ALIASES:
                continue
            if current:
                entries.append(current)
            current = {"name": title, "tech": None, "bullets": []}

        # ---- Italic-only line: *Role | tech*  (tech annotation without a name) ----
        elif (
            stripped.startswith("*")
            and stripped.endswith("*")
            and len(stripped) > 2
            and not stripped.startswith("**")
        ):
            if current is not None:
                current["tech"] = _strip_md(stripped)

        # ---- Bullet point: - text  or  * text (but not bold ** lines) ----
        elif stripped.startswith(("-",)) or (stripped.startswith("*") and not stripped.startswith("**")):
            if current is None:
                current = {"name": "", "tech": None, "bullets": []}
            bullet_text = stripped.lstrip("-* ").strip()
            if bullet_text:
                current["bullets"].append(bullet_text)

    if current:
        entries.append(current)

    return entries


# ---------------------------------------------------------------------------
# PDF helpers – font detection
# ---------------------------------------------------------------------------

# Map common resume font names (lower-cased, spaces/dashes stripped) to
# built-in reportlab PDF font families.
_PDF_FONT_MAP: dict[str, str] = {
    "helvetica": "Helvetica",
    "helveticaneue": "Helvetica",
    "arial": "Helvetica",
    "arialmt": "Helvetica",
    "calibri": "Helvetica",
    "calibribody": "Helvetica",
    "gill": "Helvetica",
    "gillsans": "Helvetica",
    "verdana": "Helvetica",
    "tahoma": "Helvetica",
    "trebuchet": "Helvetica",
    "futura": "Helvetica",
    "lato": "Helvetica",
    "opensans": "Helvetica",
    "roboto": "Helvetica",
    "sourcesans": "Helvetica",
    "times": "Times-Roman",
    "timesnewroman": "Times-Roman",
    "timesnewromanps": "Times-Roman",
    "georgia": "Times-Roman",
    "cambria": "Times-Roman",
    "garamond": "Times-Roman",
    "palatino": "Times-Roman",
    "bookman": "Times-Roman",
    "courier": "Courier",
    "couriernew": "Courier",
    "lucidaconsole": "Courier",
}


def _normalise_font_key(raw: str) -> str:
    """Strip common suffixes/prefixes and lower-case a PDF BaseFont name."""
    name = raw.lstrip("/")
    # Remove subset prefix like "ABCDEF+"
    if "+" in name:
        name = name.split("+", 1)[1]
    # Drop style suffixes: -Bold, -Italic, ,Bold, MT, PS, MSSans …
    name = re.sub(r"[-,](Bold|Italic|Regular|Light|Medium|Oblique|Roman|MT|PS).*", "", name, flags=re.IGNORECASE)
    name = re.sub(r"(Bold|Italic|Regular|Light|Medium|Oblique|Roman|MT|PS)$", "", name, flags=re.IGNORECASE)
    return name.lower().replace(" ", "").replace("-", "")


def _get_pdf_font_info(pdf_bytes: bytes) -> Tuple[str, str, str]:
    """Return ``(base_font, bold_font, italic_font)`` for the reportlab render.

    Inspects the resource dictionary of the first page to find the dominant font
    and maps it to the nearest built-in reportlab family.  Falls back to Helvetica.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(pdf_bytes))
        page = reader.pages[0]
        resources = page.get("/Resources")
        if not resources:
            raise ValueError("no resources")
        fonts_dict = resources.get("/Font")
        if not fonts_dict:
            raise ValueError("no fonts")

        # Collect BaseFont values
        seen: dict[str, int] = {}
        for _ref_name, font_ref in fonts_dict.items():
            try:
                font_obj = font_ref.get_object() if hasattr(font_ref, "get_object") else font_ref
                base_font = str(font_obj.get("/BaseFont", ""))
                if base_font:
                    key = _normalise_font_key(base_font)
                    seen[key] = seen.get(key, 0) + 1
            except Exception:
                pass

        # Pick the most frequent font
        if seen:
            dominant = max(seen, key=lambda k: seen[k])
            family = _PDF_FONT_MAP.get(dominant, "Helvetica")
        else:
            family = "Helvetica"
    except Exception:
        family = "Helvetica"

    _bold_map = {"Helvetica": "Helvetica-Bold", "Times-Roman": "Times-Bold", "Courier": "Courier-Bold"}
    _italic_map = {"Helvetica": "Helvetica-Oblique", "Times-Roman": "Times-Italic", "Courier": "Courier-Oblique"}
    return family, _bold_map.get(family, "Helvetica-Bold"), _italic_map.get(family, "Helvetica-Oblique")


# ---------------------------------------------------------------------------
# PDF – build projects page
# ---------------------------------------------------------------------------


def _build_projects_pdf_bytes(project_entries: List[dict], pdf_bytes: Optional[bytes] = None) -> bytes:
    """Render a Projects section page as PDF bytes using reportlab.

    Font family is detected from *pdf_bytes* when provided so the added page
    visually matches the rest of the document.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    body_font, bold_font, italic_font = _get_pdf_font_info(pdf_bytes) if pdf_bytes else ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique")

    buffer = io.BytesIO()
    doc_rl = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=1.0 * inch,
        rightMargin=1.0 * inch,
    )
    styles = getSampleStyleSheet()

    project_name_style = ParagraphStyle(
        "ProjName",
        parent=styles["Normal"],
        fontName=bold_font,
        fontSize=11,
        spaceAfter=1,
        spaceBefore=8,
    )
    tech_style = ParagraphStyle(
        "ProjTech",
        parent=styles["Normal"],
        fontName=italic_font,
        fontSize=9,
        spaceAfter=1,
        textColor=colors.HexColor("#525252"),
    )
    bullet_style = ParagraphStyle(
        "ProjBullet",
        parent=styles["Normal"],
        fontName=body_font,
        fontSize=10,
        leftIndent=14,
        firstLineIndent=0,
        spaceAfter=2,
    )

    story: list = []

    for entry in project_entries:
        if entry.get("name"):
            story.append(Paragraph(entry["name"], project_name_style))
        if entry.get("tech"):
            story.append(Paragraph(entry["tech"], tech_style))
        for bullet in entry.get("bullets", []):
            story.append(Paragraph(f"\u2022\u00a0{bullet}", bullet_style))

    doc_rl.build(story)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF – detect whether a projects section already exists
# ---------------------------------------------------------------------------


def _pdf_has_projects_section(pdf_bytes: bytes) -> bool:
    """Return True if the PDF's extracted text contains a Projects heading."""
    text = extract_text_from_pdf(pdf_bytes).lower()
    for alias in _PROJECTS_ALIASES:
        # Look for the alias on its own line / after a newline
        if re.search(rf"(^|\n)\s*{re.escape(alias)}\s*(\n|$)", text):
            return True
    return False


# ---------------------------------------------------------------------------
# PDF injection (public)
# ---------------------------------------------------------------------------


def inject_projects_into_pdf(pdf_bytes: bytes, projects_markdown: str) -> bytes:
    """Append project entries as a new page to an existing PDF.

    Original pages are preserved unchanged.  The appended page is rendered by
    reportlab using the same font family as the source document.  No section
    heading is added – entries (name, tech, bullets) are placed directly.
    """
    try:
        from pypdf import PdfReader, PdfWriter

        project_entries = _parse_projects_markdown(projects_markdown)
        if not project_entries:
            return pdf_bytes

        # Detect font once so we don't repeat the work
        body_font, bold_font, italic_font = _get_pdf_font_info(pdf_bytes)

        # Build the projects page, passing the font info directly
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate

        rl_buf = io.BytesIO()
        doc_rl = SimpleDocTemplate(
            rl_buf,
            pagesize=letter,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
            leftMargin=1.0 * inch,
            rightMargin=1.0 * inch,
        )
        styles = getSampleStyleSheet()

        proj_name_s = ParagraphStyle(
            "PPN",
            parent=styles["Normal"],
            fontName=bold_font,
            fontSize=11,
            spaceAfter=1,
            spaceBefore=8,
        )
        tech_s = ParagraphStyle(
            "PT",
            parent=styles["Normal"],
            fontName=italic_font,
            fontSize=9,
            spaceAfter=1,
            textColor=colors.HexColor("#525252"),
        )
        bullet_s = ParagraphStyle(
            "PB",
            parent=styles["Normal"],
            fontName=body_font,
            fontSize=10,
            leftIndent=14,
            firstLineIndent=0,
            spaceAfter=2,
        )

        story: list = []
        for entry in project_entries:
            if entry.get("name"):
                story.append(Paragraph(entry["name"], proj_name_s))
            if entry.get("tech"):
                story.append(Paragraph(entry["tech"], tech_s))
            for bullet in entry.get("bullets", []):
                story.append(Paragraph(f"\u2022\u00a0{bullet}", bullet_s))

        doc_rl.build(story)
        new_page_bytes = rl_buf.getvalue()

        # Merge
        writer = PdfWriter()
        for page in PdfReader(io.BytesIO(pdf_bytes)).pages:
            writer.add_page(page)
        for page in PdfReader(io.BytesIO(new_page_bytes)).pages:
            writer.add_page(page)

        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()

    except Exception:
        logger.exception("Failed to inject projects into PDF; returning original")
        return pdf_bytes


# ---------------------------------------------------------------------------
# DOCX helpers – font detection
# ---------------------------------------------------------------------------


def _get_docx_font_info(doc) -> Tuple[str, float]:
    """Return ``(font_name, font_size_pt)`` from the document's Normal style.

    Falls back to ``("Calibri", 11.0)`` (Microsoft Word default) when not found.
    """
    default_font = "Calibri"
    default_size = 11.0

    try:
        normal = doc.styles["Normal"]
        fname = normal.font.name
        fsize = normal.font.size
        if fname:
            default_font = fname
        if fsize:
            default_size = fsize.pt
    except Exception:
        pass

    # If Normal style font is unset, fall back to the document defaults part
    if default_font == "Calibri":
        try:
            from docx.oxml.ns import qn

            doc_defaults = doc.element.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docDefaults"
            )
            if doc_defaults is not None:
                rpr_default = doc_defaults.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPrDefault"
                )
                if rpr_default is not None:
                    rpr = rpr_default.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
                    )
                    if rpr is not None:
                        r_fonts = rpr.find(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                        )
                        if r_fonts is not None:
                            for attr in (
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
                            ):
                                v = r_fonts.get(attr)
                                if v:
                                    default_font = v
                                    break
        except Exception:
            pass

    return default_font, default_size


# ---------------------------------------------------------------------------
# DOCX helpers – style availability
# ---------------------------------------------------------------------------


def _has_style(doc, style_name: str) -> bool:
    """Return True if the document's style table contains *style_name*."""
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def _resolve_heading_style(doc, level: int) -> Optional[str]:
    """Return the style name for a heading at *level*, or None if absent."""
    candidates = [f"Heading {level}", f"heading {level}"]
    for c in candidates:
        if _has_style(doc, c):
            return c
    return None


# ---------------------------------------------------------------------------
# DOCX helpers – paragraph creation matching document font
# ---------------------------------------------------------------------------


def _set_run_font(run, font_name: str, font_size_pt: float, bold: bool = False, italic: bool = False) -> None:
    """Apply *font_name* / size / style to a run's font object."""
    from docx.shared import Pt

    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    run.italic = italic


def _add_heading_para(doc, text: str, level: int, font_name: str, font_size_pt: float):
    """Add a heading paragraph using the document's native heading style when available,
    otherwise fall back to a bold Normal paragraph with a slightly larger font."""
    style_name = _resolve_heading_style(doc, level)
    if style_name:
        para = doc.add_heading(text, level=level)
        # Override font on all runs to match the document body font
        for run in para.runs:
            run.font.name = font_name
    else:
        para = doc.add_paragraph()
        run = para.add_run(text)
        _set_run_font(run, font_name, font_size_pt + (3 - level), bold=True)
    return para


def _add_bullet_para(doc, text: str, font_name: str, font_size_pt: float):
    """Add a bullet paragraph with font matching the document."""
    if _has_style(doc, "List Bullet"):
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        _set_run_font(run, font_name, font_size_pt)
    else:
        from docx.shared import Pt

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        run = para.add_run(f"\u2022\u00a0{text}")
        _set_run_font(run, font_name, font_size_pt)
    return para


def _add_italic_para(doc, text: str, font_name: str, font_size_pt: float):
    """Add an italicised paragraph matching the document font."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, font_name, font_size_pt - 1, italic=True)
    return para


# ---------------------------------------------------------------------------
# DOCX helpers – locate existing Projects section
# ---------------------------------------------------------------------------


def _find_projects_section(paragraphs) -> Tuple[Optional[int], int]:
    """Return ``(heading_para_idx, heading_level)`` for the Projects section.

    Returns ``(None, 2)`` when no Projects heading is found.
    """
    for i, para in enumerate(paragraphs):
        text_lower = para.text.strip().lower()
        if text_lower not in _PROJECTS_ALIASES:
            continue
        style_name = (para.style.name or "").lower()
        is_heading = "heading" in style_name
        is_bold_line = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
        if is_heading or is_bold_line:
            level = 2
            if is_heading:
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 2
            return i, level
    return None, 2


def _find_section_end(paragraphs, start_idx: int, heading_level: int) -> int:
    """Return the index of the last paragraph belonging to the section starting at *start_idx*."""
    end_idx = start_idx
    for i in range(start_idx + 1, len(paragraphs)):
        style_name = (paragraphs[i].style.name or "").lower()
        if "heading" in style_name:
            try:
                lvl = int(style_name.split()[-1])
                if lvl <= heading_level:
                    break
            except (ValueError, IndexError):
                break
        end_idx = i
    return end_idx


# ---------------------------------------------------------------------------
# DOCX injection – XML-level insertion so we can place runs mid-document
# ---------------------------------------------------------------------------


def _make_xml_paragraph(
    text: str,
    font_name: str,
    font_size_pt: float,
    bold: bool = False,
    italic: bool = False,
    style_id: Optional[str] = None,
    bullet_indent: bool = False,
):
    """Build a raw ``w:p`` XML element with optional style, font, and formatting."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _tag(name: str):
        return f"{{{WNS}}}{name}"

    p = OxmlElement("w:p")

    # --- paragraph properties ---
    pPr = OxmlElement("w:pPr")
    if style_id:
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)
    if bullet_indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:hanging"), "180")
        pPr.append(ind)
    p.append(pPr)

    # --- run ---
    r = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    # Size (in half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(font_size_pt * 2)))
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)

    return p


# ---------------------------------------------------------------------------
# DOCX injection (public)
# ---------------------------------------------------------------------------


def inject_projects_into_docx(docx_bytes: bytes, projects_markdown: str) -> bytes:
    """Inject project entries into an existing DOCX resume.

    * If a Projects section heading already exists (matching ``_PROJECTS_ALIASES``),
      new project names and their resume bullets are inserted at the end of that
      section using XML-level DOM manipulation so the font matches the document.
    * If no Projects section exists, one is appended at the end of the document.

    All added text uses the same font family and approximate size as the
    document's Normal style so the output blends with the original layout.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        project_entries = _parse_projects_markdown(projects_markdown)

        if not project_entries:
            out = io.BytesIO()
            doc.save(out)
            return out.getvalue()

        font_name, font_size = _get_docx_font_info(doc)
        paragraphs = doc.paragraphs

        proj_idx, proj_level = _find_projects_section(paragraphs)

        if proj_idx is None:
            # ---- No Projects section found – append entries directly at end ----
            for entry in project_entries:
                if entry.get("name"):
                    _add_heading_para(doc, entry["name"], 3, font_name, font_size)
                if entry.get("tech"):
                    _add_italic_para(doc, entry["tech"], font_name, font_size)
                for bullet in entry.get("bullets", []):
                    _add_bullet_para(doc, bullet, font_name, font_size)
        else:
            # ---- Insert at end of existing Projects section ----
            end_idx = _find_section_end(paragraphs, proj_idx, proj_level)
            ref_elem = paragraphs[end_idx]._element

            # Build and attach paragraphs in reverse order via addnext
            # so they appear in the correct sequence after ref_elem.
            all_new: list = []
            for entry in project_entries:
                if entry.get("name"):
                    all_new.append(
                        _make_xml_paragraph(
                            entry["name"],
                            font_name,
                            font_size + 1,
                            bold=True,
                        )
                    )
                if entry.get("tech"):
                    all_new.append(
                        _make_xml_paragraph(
                            entry["tech"],
                            font_name,
                            font_size - 1,
                            italic=True,
                        )
                    )
                for bullet in entry.get("bullets", []):
                    bullet_text = f"\u2022\u00a0{bullet}"
                    all_new.append(
                        _make_xml_paragraph(
                            bullet_text,
                            font_name,
                            font_size,
                            bullet_indent=True,
                        )
                    )

            # Insert in order: addnext inserts immediately after ref_elem each time
            for xml_para in all_new:
                ref_elem.addnext(xml_para)
                ref_elem = xml_para

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    except Exception:
        logger.exception("Failed to inject projects into DOCX; returning original")
        return docx_bytes
